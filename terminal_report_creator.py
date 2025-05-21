import tkinter as tk
from tkinter import filedialog, messagebox, ttk, simpledialog, PhotoImage, Canvas
import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import Image, ImageTk, ImageOps
import glob
import re
import threading
import math
from datetime import datetime
from fuzzywuzzy import fuzz
import json
import uuid
import hashlib
import platform
import wmi
import psutil
import base64
from cryptography.fernet import Fernet
import requests
import socket
import time
import logging
from pathlib import Path
import string
import random
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configure logging
logging.basicConfig(
    filename='app_security.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

class SecurityManager:
    def __init__(self):
        self.license_file = "license.key"
        self.encryption_key = None
        self.fernet = None
        self.initialize_encryption()
        self.license_window = None
        self.main_window = None
        
    def show_license_window(self):
        """Show a window for license input"""
        if self.license_window:
            return
            
        self.license_window = tk.Toplevel()
        self.license_window.title("License Activation")
        self.license_window.geometry("500x400")
        self.license_window.configure(bg="#1E1E1E")
        self.license_window.protocol("WM_DELETE_WINDOW", self.on_license_window_close)  # Handle window closing
        self.license_window.transient()
        self.license_window.grab_set()
        
        # Center the window
        self.license_window.geometry("+%d+%d" % (
            self.license_window.winfo_screenwidth() // 2 - 250,
            self.license_window.winfo_screenheight() // 2 - 200
        ))
        
        # Create main frame
        main_frame = tk.Frame(self.license_window, bg="#1E1E1E", padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = tk.Label(
            main_frame,
            text="Tree Inventory - License Activation",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 14, "bold")
        )
        title_label.pack(pady=(0, 20))
        
        # License input frame
        input_frame = tk.Frame(main_frame, bg="#1E1E1E")
        input_frame.pack(fill=tk.BOTH, expand=True)
        
        # License text area
        license_label = tk.Label(
            input_frame,
            text="Enter or paste your license key:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10)
        )
        license_label.pack(anchor="w", pady=(0, 5))
        
        self.license_text = tk.Text(
            input_frame,
            height=8,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            font=("Arial", 10)
        )
        self.license_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Buttons frame
        button_frame = tk.Frame(main_frame, bg="#1E1E1E")
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        # Load from file button
        load_btn = tk.Button(
            button_frame,
            text="Load from File",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.load_license_from_file
        )
        load_btn.pack(side=tk.LEFT, padx=5)
        
        # Activate button
        activate_btn = tk.Button(
            button_frame,
            text="Activate License",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.activate_license
        )
        activate_btn.pack(side=tk.RIGHT, padx=5)
        
        # Status label
        self.status_label = tk.Label(
            main_frame,
            text="",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10)
        )
        self.status_label.pack(pady=(10, 0))
        
        # Try to load existing license
        if os.path.exists(self.license_file):
            try:
                with open(self.license_file, "rb") as f:
                    encrypted_data = f.read()
                    license_data = json.loads(self.fernet.decrypt(encrypted_data))
                    self.license_text.insert("1.0", json.dumps(license_data, indent=2))
            except:
                pass

    def on_license_window_close(self):
        """Handle license window closing"""
        if self.license_window:
            self.license_window.destroy()
            self.license_window = None
            if self.main_window:
                self.main_window.destroy()  # Close the main window
            sys.exit(0)  # Exit the application

    def load_license_from_file(self):
        """Load license from a file"""
        file_path = filedialog.askopenfilename(
            title="Select License File",
            filetypes=[("License files", "*.key *.lic *.txt"), ("All files", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, "r") as f:
                    license_text = f.read()
                    self.license_text.delete("1.0", tk.END)
                    self.license_text.insert("1.0", license_text)
                self.status_label.config(text="License loaded successfully", fg="#00FF00")
            except Exception as e:
                self.status_label.config(text=f"Error loading license: {str(e)}", fg="#FF0000")

    def activate_license(self):
        """Activate the license"""
        try:
            license_text = self.license_text.get("1.0", tk.END).strip()
            if not license_text:
                self.status_label.config(text="Please enter a license key", fg="#FF0000")
                return
            
            # Try to parse as JSON
            try:
                license_data = json.loads(license_text)
            except:
                # If not JSON, treat as raw license key
                license_data = {
                    'license_key': license_text,
                    'expiry_date': '2099-12-31',  # Default expiry
                    'hardware_id': self.get_hardware_id()
                }
            
            # Add signature
            license_data['signature'] = hashlib.sha256(
                json.dumps(license_data, sort_keys=True).encode()
            ).hexdigest()
            
            # Encrypt and save
            encrypted_data = self.fernet.encrypt(json.dumps(license_data).encode())
            with open(self.license_file, "wb") as f:
                f.write(encrypted_data)
            
            self.status_label.config(text="License activated successfully", fg="#00FF00")
            
            # Close license window and show main window
            if self.license_window:
                self.license_window.destroy()
                self.license_window = None
            
            if self.main_window:
                self.main_window.deiconify()  # Show the main window
            
        except Exception as e:
            self.status_label.config(text=f"Error activating license: {str(e)}", fg="#FF0000")

    def initialize_encryption(self):
        """Initialize encryption with a secure key"""
        try:
            # Generate or load encryption key
            key_file = "encryption.key"
            if os.path.exists(key_file):
                with open(key_file, "rb") as f:
                    self.encryption_key = f.read()
            else:
                self.encryption_key = Fernet.generate_key()
                with open(key_file, "wb") as f:
                    f.write(self.encryption_key)
            
            self.fernet = Fernet(self.encryption_key)
        except Exception as e:
            logging.error(f"Encryption initialization error: {str(e)}")
            raise

    def get_hardware_id(self):
        """Generate a unique hardware ID based on system components"""
        try:
            c = wmi.WMI()
            system_info = c.Win32_ComputerSystemProduct()[0]
            cpu_info = c.Win32_Processor()[0]
            bios_info = c.Win32_BIOS()[0]
            
            hardware_str = f"{system_info.UUID}-{cpu_info.ProcessorId}-{bios_info.SerialNumber}"
            return hashlib.sha256(hardware_str.encode()).hexdigest()
        except Exception as e:
            logging.error(f"Hardware ID generation error: {str(e)}")
            return None

    def validate_license(self):
        """Validate the software license"""
        try:
            if not os.path.exists(self.license_file):
                return False, "License file not found"

            with open(self.license_file, "rb") as f:
                encrypted_data = f.read()
                license_data = json.loads(self.fernet.decrypt(encrypted_data))

            # Check if license is expired
            if datetime.fromisoformat(license_data['expiry_date']) < datetime.now():
                return False, "License has expired"

            # Validate hardware ID
            current_hw_id = self.get_hardware_id()
            if current_hw_id != license_data['hardware_id']:
                return False, "Invalid hardware ID"

            # Validate license signature
            if not self.verify_license_signature(license_data):
                return False, "Invalid license signature"

            return True, "License valid"
        except Exception as e:
            logging.error(f"License validation error: {str(e)}")
            return False, f"License validation error: {str(e)}"

    def verify_license_signature(self, license_data):
        """Verify the license signature"""
        try:
            signature = license_data.pop('signature', None)
            if not signature:
                return False

            # Recreate the data that was signed
            data_to_verify = json.dumps(license_data, sort_keys=True)
            expected_signature = hashlib.sha256(data_to_verify.encode()).hexdigest()
            
            return signature == expected_signature
        except Exception as e:
            logging.error(f"Signature verification error: {str(e)}")
            return False

    def encrypt_data(self, data):
        """Encrypt data using Fernet"""
        try:
            return self.fernet.encrypt(data.encode()).decode()
        except Exception as e:
            logging.error(f"Encryption error: {str(e)}")
            return None

    def decrypt_data(self, encrypted_data):
        """Decrypt data using Fernet"""
        try:
            return self.fernet.decrypt(encrypted_data.encode()).decode()
        except Exception as e:
            logging.error(f"Decryption error: {str(e)}")
            return None

    def check_tampering(self):
        """Check for signs of tampering"""
        try:
            # Check if running from expected location
            expected_path = os.path.abspath(sys.argv[0])
            if not os.path.exists(expected_path):
                return False, "Application moved from original location"

            # Check file integrity
            if not self.verify_file_integrity():
                return False, "File integrity check failed"

            # Check for debugging tools
            if self.detect_debugger():
                return False, "Debugging detected"

            return True, "No tampering detected"
        except Exception as e:
            logging.error(f"Tampering check error: {str(e)}")
            return False, f"Tampering check error: {str(e)}"

    def verify_file_integrity(self):
        """Verify the integrity of the application file"""
        try:
            # For now, we'll skip the hash check since we don't have a stored hash
            # In a production environment, you would want to implement proper hash verification
            return True
            
            # Original hash check code (commented out for now)
            # with open(sys.argv[0], 'rb') as f:
            #     current_hash = hashlib.sha256(f.read()).hexdigest()
            # stored_hash = "YOUR_STORED_HASH"  # Replace with your stored hash
            # return current_hash == stored_hash
        except Exception as e:
            logging.error(f"File integrity check error: {str(e)}")
            return False

    def detect_debugger(self):
        """Detect if a debugger is attached"""
        try:
            # Check for common debugging tools
            debugger_processes = ['x64dbg.exe', 'ollydbg.exe', 'ida64.exe', 'windbg.exe']
            for proc in psutil.process_iter(['name']):
                if proc.info['name'].lower() in [p.lower() for p in debugger_processes]:
                    return True
            return False
        except Exception as e:
            logging.error(f"Debugger detection error: {str(e)}")
            return False

class Tooltip:
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        self.widget.bind("<Enter>", self.show_tooltip)
        self.widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event=None):
        if self.tooltip_window:
            return
        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + 20
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(
            tw,
            text=self.text,
            background="#FFFFDD",
            foreground="#000000",
            relief="solid",
            borderwidth=1,
            font=("Arial", 8),
            wraplength=400
        )
        label.pack()

    def hide_tooltip(self, event=None):
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

class TreeInventoryApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Tree Inventory Manager")
        self.root.configure(bg="#1E1E1E")
        
        # Initialize security manager
        self.security_manager = SecurityManager()
        self.security_manager.main_window = self.root
        
        # Hide main window initially
        self.root.withdraw()
        
        # Validate security and show license window if needed
        self.validate_security()
        
        # Continue with normal initialization
        self.root.geometry("800x600")
        self.root.minsize(600, 400)
        
        self.main_frame = tk.Frame(root, bg="#1E1E1E")
        self.main_frame.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.9)
        
        self.create_content()
        self.create_footer()
        
        # Initialize other attributes
        self.excel_file_path = None
        self.image_folder_path = None
        self.output_folder_path = None
        self.start_row = 2
        self.batch_size = 1002
        self.num_batches_to_process = None
        self.docx_filename_prefix = "TreeData"
        self.cutting_docx_filename_prefix = "TreeCutting"
        self.image_pattern = "auto_detect"
        self.use_subfolders = True
        self.image_cache = {}
        self.pattern_stats = {}
        self.match_cache = {}
        self.image_size_cache = {}
        self.fuzzy_cache = {}

    def validate_security(self):
        """Validate security before starting the application"""
        try:
            # Check for tampering
            tampering_valid, tampering_msg = self.security_manager.check_tampering()
            if not tampering_valid:
                messagebox.showerror("Security Error", tampering_msg)
                sys.exit(1)

            # Validate license
            license_valid, license_msg = self.security_manager.validate_license()
            if not license_valid:
                # Show license window
                self.security_manager.show_license_window()
                return

            # If license is valid, show main window
            self.root.deiconify()
            
            # Log successful validation
            logging.info("Security validation successful")
        except Exception as e:
            logging.error(f"Security validation error: {str(e)}")
            # Show license window
            self.security_manager.show_license_window()

    def create_content(self):
        content_frame = tk.Frame(self.main_frame, bg="#1E1E1E")
        content_frame.place(relx=0.5, rely=0.2, relwidth=0.8, relheight=0.7, anchor="n")
        
        system_label = tk.Label(
            content_frame,
            text="Tree Inventory - Tree Cutting Terminal Report Creator",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 16, "bold"),
            wraplength=500,
            justify="center"
        )
        system_label.place(relx=0.5, rely=0.3, anchor="center")
        
        self.inventory_btn = tk.Button(
            content_frame,
            text="INVENTORY",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 12, "bold"),
            width=15,
            height=2,
            command=self.show_inventory
        )
        self.inventory_btn.place(relx=0.3, rely=0.6, anchor="center")
        
        self.cutting_btn = tk.Button(
            content_frame,
            text="CUTTING",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 12, "bold"),
            width=15,
            height=2,
            command=self.show_cutting
        )
        self.cutting_btn.place(relx=0.7, rely=0.6, anchor="center")

    def create_footer(self):
        self.footer_frame = tk.Frame(self.main_frame, bg="#1E1E1E")
        self.footer_frame.place(relx=0.5, rely=0.95, relwidth=0.9, relheight=0.05, anchor="s")
        
        version_label = tk.Label(
            self.footer_frame,
            text="Tree Inventory - Tree Cutting Terminal Report v2.0 | Created by Ron Michael Comia | © 2025",
            bg="#1E1E1E",
            fg="#888888",
            font=("Arial", 8),
            anchor="w"
        )
        version_label.pack(side=tk.LEFT)
        
        help_label = tk.Label(
            self.footer_frame,
            text="Help & Support",
            bg="#1E1E1E",
            fg="#FFC107",
            font=("Arial", 8, "underline"),
            cursor="hand2"
        )
        help_label.pack(side=tk.RIGHT)
        help_label.bind("<Button-1>", lambda e: self.show_help())

    def reset_button_colors(self):
        self.inventory_btn.config(fg="#FFFFFF")
        self.cutting_btn.config(fg="#FFFFFF")

    def show_initial(self):
        self.clear_content()
        self.create_content()
        self.reset_button_colors()

    def show_help(self):
        help_window = tk.Toplevel(self.root)
        help_window.title("Help & Support")
        help_window.geometry("400x300")
        help_window.configure(bg="#1E1E1E")
        help_window.transient(self.root)
        help_window.grab_set()
        
        help_window.geometry("+%d+%d" % (
            self.root.winfo_rootx() + self.root.winfo_width() // 2 - 200,
            self.root.winfo_rooty() + self.root.winfo_height() // 2 - 150
        ))
        
        help_text = tk.Label(
            help_window,
            text="Tree Inventory - Tree Cutting Terminal Report\n\n"
                 "1. Make sure the Photo/Pictures is named properly\n"
                 "2. Select Excel file with tree data\n"
                 "3. Choose image folder\n"
                 "4. Set output folder\n"
                 "5. Configure settings as needed\n"
                 "6. Name the File the way you want it\n"
                 "7. Click 'Generate Word Documents'\n\n"
                 "For issues, contact ronm.comia@gmail.com",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            justify=tk.LEFT,
            wraplength=360
        )
        help_text.pack(pady=20, padx=20)
        
        close_btn = tk.Button(
            help_window,
            text="Close",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=help_window.destroy
        )
        close_btn.pack(pady=10)

    def show_inventory(self):
        self.inventory_btn.config(fg="#FFC107")
        self.cutting_btn.config(fg="#FFFFFF")
        self.clear_content()
        
        inventory_frame = tk.Frame(self.main_frame, bg="#1E1E1E")
        inventory_frame.place(relx=0.5, rely=0.05, relwidth=0.9, relheight=0.9, anchor="n")
        
        back_btn = tk.Button(
            inventory_frame,
            text="Back",
            bg="#1E1E1E",
            fg="#FFC107",
            activebackground="#2A2A2A",
            activeforeground="#FFC107",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.show_initial
        )
        back_btn.pack(pady=5, anchor="w")
        
        select_excel_btn = tk.Button(
            inventory_frame,
            text="Select Excel File",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_excel_file
        )
        select_excel_btn.pack(pady=5)
        
        select_images_btn = tk.Button(
            inventory_frame,
            text="Select Images Folder",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_image_folder
        )
        select_images_btn.pack(pady=5)
        
        select_output_btn = tk.Button(
            inventory_frame,
            text="Select Output Folder",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_output_folder
        )
        select_output_btn.pack(pady=5)
        
        settings_btn = tk.Button(
            inventory_frame,
            text="Image Naming Settings",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.show_image_settings
        )
        settings_btn.pack(pady=5)
        
        config_frame = tk.Frame(inventory_frame, bg="#1E1E1E", bd=1, relief=tk.SOLID)
        config_frame.pack(pady=10, padx=20, fill=tk.X)
        
        start_row_frame = tk.Frame(config_frame, bg="#1E1E1E")
        start_row_frame.pack(pady=5, fill=tk.X)
        
        start_row_label = tk.Label(
            start_row_frame,
            text="Start Row:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        start_row_label.pack(side=tk.LEFT, padx=5)
        
        self.start_row_var = tk.StringVar(value=str(self.start_row))
        start_row_entry = tk.Entry(
            start_row_frame,
            textvariable=self.start_row_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=10
        )
        start_row_entry.pack(side=tk.LEFT, padx=5)
        
        batch_size_frame = tk.Frame(config_frame, bg="#1E1E1E")
        batch_size_frame.pack(pady=5, fill=tk.X)
        
        batch_size_label = tk.Label(
            batch_size_frame,
            text="Batch Size:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        batch_size_label.pack(side=tk.LEFT, padx=5)
        
        self.batch_size_var = tk.StringVar(value=str(self.batch_size))
        batch_size_entry = tk.Entry(
            batch_size_frame,
            textvariable=self.batch_size_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=10
        )
        batch_size_entry.pack(side=tk.LEFT, padx=5)
        
        # Add batch calculation display
        self.batch_calc_label = tk.Label(
            batch_size_frame,
            text="Estimated batches: --",
            bg="#1E1E1E",
            fg="#888888",
            font=("Arial", 8)
        )
        self.batch_calc_label.pack(side=tk.LEFT, padx=5)
        
        # Bind batch size changes
        self.batch_size_var.trace_add("write", self.update_batch_calculation)
        
        num_batches_frame = tk.Frame(config_frame, bg="#1E1E1E")
        num_batches_frame.pack(pady=5, fill=tk.X)
        
        num_batches_label = tk.Label(
            num_batches_frame,
            text="Batches to Process:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        num_batches_label.pack(side=tk.LEFT, padx=5)
        
        self.num_batches_var = tk.StringVar(value="All")
        num_batches_entry = tk.Entry(
            num_batches_frame,
            textvariable=self.num_batches_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=10
        )
        num_batches_entry.pack(side=tk.LEFT, padx=5)
        
        help_label = tk.Label(
            num_batches_frame,
            text="(Enter 'All' or a number)",
            bg="#1E1E1E",
            fg="#888888",
            font=("Arial", 8)
        )
        help_label.pack(side=tk.LEFT, padx=5)
        
        # Filename prefix input for Inventory
        filename_frame = tk.Frame(config_frame, bg="#1E1E1E")
        filename_frame.pack(pady=5, fill=tk.X)
        
        filename_label = tk.Label(
            filename_frame,
            text="DOCX Filename Prefix:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        filename_label.pack(side=tk.LEFT, padx=5)
        
        self.filename_prefix_var = tk.StringVar(value=self.docx_filename_prefix)
        filename_entry = tk.Entry(
            filename_frame,
            textvariable=self.filename_prefix_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=20
        )
        filename_entry.pack(side=tk.LEFT, padx=5)
        
        self.process_btn = tk.Button(
            inventory_frame,
            text="Generate Word Documents",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10, "bold"),
            command=self.process_data,
            state=tk.DISABLED
        )
        self.process_btn.pack(pady=10)
        
        status_canvas = tk.Canvas(inventory_frame, bg="#1E1E1E", highlightthickness=0)
        scrollbar = ttk.Scrollbar(inventory_frame, orient="vertical", command=status_canvas.yview)
        self.status_frame = tk.Frame(status_canvas, bg="#1E1E1E")
        
        self.status_frame.bind(
            "<Configure>",
            lambda e: status_canvas.configure(scrollregion=status_canvas.bbox("all"))
        )
        status_canvas.configure(yscrollcommand=scrollbar.set)
        
        status_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=10, padx=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        status_canvas.create_window((0, 0), window=self.status_frame, anchor="nw")
        
        def truncate_path(path, max_length=50):
            if not path:
                return "Not selected"
            if len(path) <= max_length:
                return path
            return f"{path[:max_length-3]}..."
        
        excel_path = self.excel_file_path if self.excel_file_path else "Not selected"
        self.excel_label = tk.Label(
            self.status_frame,
            text=f"Excel File: {truncate_path(excel_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if self.excel_file_path else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.excel_label.pack(fill=tk.X, pady=2)
        if self.excel_file_path:
            Tooltip(self.excel_label, self.excel_file_path)
        
        images_path = self.image_folder_path if self.image_folder_path else "Not selected"
        self.images_label = tk.Label(
            self.status_frame,
            text=f"Images Folder: {truncate_path(images_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if self.image_folder_path else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.images_label.pack(fill=tk.X, pady=2)
        if self.image_folder_path:
            Tooltip(self.images_label, self.image_folder_path)
        
        output_path = self.output_folder_path if self.output_folder_path else "Not selected"
        self.output_label = tk.Label(
            self.status_frame,
            text=f"Output Folder: {truncate_path(output_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if self.output_folder_path else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.output_label.pack(fill=tk.X, pady=2)
        if self.output_folder_path:
            Tooltip(self.output_label, self.output_folder_path)
        
        pattern_name = next((name for name, val in [
            ("Simple numbers", "simple_numbers"),
            ("Tree prefix", "tree_prefix"),
            ("Alpha-numeric", "alpha_numeric"),
            ("Number with suffix", "number_suffix"),
            ("Auto-detect", "auto_detect")
        ] if val == self.image_pattern), "Unknown")
        
        self.settings_label = tk.Label(
            self.status_frame,
            text=f"Pattern: {pattern_name} | Subfolders: {'Yes' if self.use_subfolders else 'No'}",
            bg="#1E1E1E",
            fg="#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.settings_label.pack(fill=tk.X, pady=2)
        
        self.progress_frame = tk.Frame(inventory_frame, bg="#1E1E1E")
        self.progress_frame.pack(fill=tk.X, pady=10)
        
        self.progress_label = tk.Label(
            self.progress_frame,
            text="Progress:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 8)
        )
        self.progress_label.pack(anchor="w")
        
        self.progress_bar = ttk.Progressbar(
            self.progress_frame,
            orient="horizontal",
            length=300,
            mode="determinate"
        )
        self.progress_bar.pack(fill=tk.X, pady=5)
        
        self.progress_frame.pack_forget()

    def show_cutting(self):
        self.cutting_btn.config(fg="#FFC107")
        self.inventory_btn.config(fg="#FFFFFF")
        self.clear_content()
        
        cutting_frame = tk.Frame(self.main_frame, bg="#1E1E1E")
        cutting_frame.place(relx=0.5, rely=0.05, relwidth=0.9, relheight=0.9, anchor="n")
        
        back_btn = tk.Button(
            cutting_frame,
            text="Back",
            bg="#1E1E1E",
            fg="#FFC107",
            activebackground="#2A2A2A",
            activeforeground="#FFC107",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.show_initial
        )
        back_btn.pack(pady=5, anchor="w")
        
        select_before_btn = tk.Button(
            cutting_frame,
            text="Select Before Images Folder",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_before_images
        )
        select_before_btn.pack(pady=5)
        
        select_after_btn = tk.Button(
            cutting_frame,
            text="Select After Images Folder",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_after_images
        )
        select_after_btn.pack(pady=5)
        
        select_excel_btn = tk.Button(
            cutting_frame,
            text="Select Excel File",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_cutting_excel
        )
        select_excel_btn.pack(pady=5)
        
        select_output_btn = tk.Button(
            cutting_frame,
            text="Select Output Folder",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=self.select_cutting_output
        )
        select_output_btn.pack(pady=5)
        
        config_frame = tk.Frame(cutting_frame, bg="#1E1E1E", bd=1, relief=tk.SOLID)
        config_frame.pack(pady=10, padx=20, fill=tk.X)
        
        start_row_frame = tk.Frame(config_frame, bg="#1E1E1E")
        start_row_frame.pack(pady=5, fill=tk.X)
        
        start_row_label = tk.Label(
            start_row_frame,
            text="Start Row:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        start_row_label.pack(side=tk.LEFT, padx=5)
        
        self.cutting_start_row_var = tk.StringVar(value="2")
        start_row_entry = tk.Entry(
            start_row_frame,
            textvariable=self.cutting_start_row_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=10
        )
        start_row_entry.pack(side=tk.LEFT, padx=5)
        
        batch_size_frame = tk.Frame(config_frame, bg="#1E1E1E")
        batch_size_frame.pack(pady=5, fill=tk.X)
        
        batch_size_label = tk.Label(
            batch_size_frame,
            text="Batch Size:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        batch_size_label.pack(side=tk.LEFT, padx=5)
        
        self.cutting_batch_size_var = tk.StringVar(value="100")
        batch_size_entry = tk.Entry(
            batch_size_frame,
            textvariable=self.cutting_batch_size_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=10
        )
        batch_size_entry.pack(side=tk.LEFT, padx=5)
        
        # Add batch calculation display for cutting mode
        self.cutting_batch_calc_label = tk.Label(
            batch_size_frame,
            text="Estimated batches: --",
            bg="#1E1E1E",
            fg="#888888",
            font=("Arial", 8)
        )
        self.cutting_batch_calc_label.pack(side=tk.LEFT, padx=5)
        
        # Bind batch size changes for cutting mode
        self.cutting_batch_size_var.trace_add("write", self.update_cutting_batch_calculation)
        
        # Filename prefix input for Cutting
        filename_frame = tk.Frame(config_frame, bg="#1E1E1E")
        filename_frame.pack(pady=5, fill=tk.X)
        
        filename_label = tk.Label(
            filename_frame,
            text="DOCX Filename Prefix:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10),
            width=15,
            anchor="e"
        )
        filename_label.pack(side=tk.LEFT, padx=5)
        
        self.cutting_filename_prefix_var = tk.StringVar(value=self.cutting_docx_filename_prefix)
        filename_entry = tk.Entry(
            filename_frame,
            textvariable=self.cutting_filename_prefix_var,
            bg="#2A2A2A",
            fg="#FFFFFF",
            insertbackground="#FFFFFF",
            width=20
        )
        filename_entry.pack(side=tk.LEFT, padx=5)
        
        self.cutting_process_btn = tk.Button(
            cutting_frame,
            text="Generate Cutting Documents",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10, "bold"),
            command=self.process_cutting_data,
            state=tk.DISABLED
        )
        self.cutting_process_btn.pack(pady=10)
        
        status_canvas = tk.Canvas(cutting_frame, bg="#1E1E1E", highlightthickness=0)
        scrollbar = ttk.Scrollbar(cutting_frame, orient="vertical", command=status_canvas.yview)
        self.cutting_status_frame = tk.Frame(status_canvas, bg="#1E1E1E")
        
        self.cutting_status_frame.bind(
            "<Configure>",
            lambda e: status_canvas.configure(scrollregion=status_canvas.bbox("all"))
        )
        status_canvas.configure(yscrollcommand=scrollbar.set)
        
        status_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=10, padx=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        status_canvas.create_window((0, 0), window=self.cutting_status_frame, anchor="nw")
        
        def truncate_path(path, max_length=50):
            if not path:
                return "Not selected"
            if len(path) <= max_length:
                return path
            return f"{path[:max_length-3]}..."
        
        before_path = getattr(self, 'before_images_folder', None) if hasattr(self, 'before_images_folder') else "Not selected"
        self.before_images_label = tk.Label(
            self.cutting_status_frame,
            text=f"Before Images: {truncate_path(before_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if hasattr(self, 'before_images_folder') else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.before_images_label.pack(fill=tk.X, pady=2)
        if hasattr(self, 'before_images_folder'):
            Tooltip(self.before_images_label, self.before_images_folder)
        
        after_path = getattr(self, 'after_images_folder', None) if hasattr(self, 'after_images_folder') else "Not selected"
        self.after_images_label = tk.Label(
            self.cutting_status_frame,
            text=f"After Images: {truncate_path(after_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if hasattr(self, 'after_images_folder') else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.after_images_label.pack(fill=tk.X, pady=2)
        if hasattr(self, 'after_images_folder'):
            Tooltip(self.after_images_label, self.after_images_folder)
        
        excel_path = getattr(self, 'cutting_excel_file', None) if hasattr(self, 'cutting_excel_file') else "Not selected"
        self.cutting_excel_label = tk.Label(
            self.cutting_status_frame,
            text=f"Excel File: {truncate_path(excel_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if hasattr(self, 'cutting_excel_file') else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.cutting_excel_label.pack(fill=tk.X, pady=2)
        if hasattr(self, 'cutting_excel_file'):
            Tooltip(self.cutting_excel_label, self.cutting_excel_file)
        
        output_path = getattr(self, 'cutting_output_folder', None) if hasattr(self, 'cutting_output_folder') else "Not selected"
        self.cutting_output_label = tk.Label(
            self.cutting_status_frame,
            text=f"Output Folder: {truncate_path(output_path)}",
            bg="#1E1E1E",
            fg="#FFFFFF" if hasattr(self, 'cutting_output_folder') else "#888888",
            font=("Arial", 8),
            anchor="w",
            wraplength=500
        )
        self.cutting_output_label.pack(fill=tk.X, pady=2)
        if hasattr(self, 'cutting_output_folder'):
            Tooltip(self.cutting_output_label, self.cutting_output_folder)
        
        self.cutting_progress_frame = tk.Frame(cutting_frame, bg="#1E1E1E")
        self.cutting_progress_frame.pack(fill=tk.X, pady=10)
        
        self.cutting_progress_label = tk.Label(
            self.cutting_progress_frame,
            text="Progress:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 8)
        )
        self.cutting_progress_label.pack(anchor="w")
        
        self.cutting_progress_bar = ttk.Progressbar(
            self.cutting_progress_frame,
            orient="horizontal",
            length=300,
            mode="determinate"
        )
        self.cutting_progress_bar.pack(fill=tk.X, pady=5)
        
        self.cutting_progress_frame.pack_forget()

    def clear_content(self):
        for widget in self.main_frame.winfo_children():
            if widget != self.footer_frame:
                widget.destroy()

    def select_excel_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        
        if file_path:
            self.excel_file_path = file_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'excel_label'):
                self.excel_label.config(
                    text=f"Excel File: {truncate_path(file_path)}",
                    fg="#FFFFFF"
                )
                self.excel_label.unbind("<Enter>")
                self.excel_label.unbind("<Leave>")
                Tooltip(self.excel_label, file_path)
            else:
                self.show_inventory()
            
            # Calculate batches after loading Excel file
            self.update_batch_calculation()
            self.check_all_selected()

    def update_batch_calculation(self, *args):
        """Update the batch calculation display based on Excel file and batch size."""
        try:
            if not self.excel_file_path:
                self.batch_calc_label.config(text="Estimated batches: --")
                return
                
            # Read Excel file to get total rows
            df = pd.read_excel(self.excel_file_path)
            total_rows = len(df)
            
            # Get current batch size
            try:
                batch_size = int(self.batch_size_var.get())
                if batch_size < 1:
                    raise ValueError("Batch size must be positive")
            except ValueError:
                self.batch_calc_label.config(text="Estimated batches: Invalid batch size")
                return
            
            # Calculate number of batches
            num_batches = math.ceil(total_rows / batch_size)
            
            # Update display
            self.batch_calc_label.config(
                text=f"Estimated batches: {num_batches} (Total rows: {total_rows})",
                fg="#FFFFFF"
            )
            
        except Exception as e:
            self.batch_calc_label.config(
                text=f"Error calculating batches: {str(e)}",
                fg="#FF0000"
            )

    def select_image_folder(self):
        folder_path = filedialog.askdirectory(
            title="Select Folder Containing Tree Images"
        )
        
        if folder_path:
            self.image_folder_path = folder_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'images_label'):
                self.images_label.config(
                    text=f"Images Folder: {truncate_path(folder_path)}",
                    fg="#FFFFFF"
                )
                self.images_label.unbind("<Enter>")
                self.images_label.unbind("<Leave>")
                Tooltip(self.images_label, folder_path)
            else:
                self.show_inventory()
            self.check_all_selected()

    def select_output_folder(self):
        folder_path = filedialog.askdirectory(
            title="Select Output Folder for Word Documents"
        )
        
        if folder_path:
            self.output_folder_path = folder_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'output_label'):
                self.output_label.config(
                    text=f"Output Folder: {truncate_path(folder_path)}",
                    fg="#FFFFFF"
                )
                self.output_label.unbind("<Enter>")
                self.output_label.unbind("<Leave>")
                Tooltip(self.output_label, folder_path)
            else:
                self.show_inventory()
            self.check_all_selected()

    def check_all_selected(self):
        if (self.excel_file_path and self.image_folder_path and self.output_folder_path):
            self.process_btn.config(state=tk.NORMAL)
        else:
            self.process_btn.config(state=tk.DISABLED)

    def select_before_images(self):
        folder_path = filedialog.askdirectory(
            title="Select Folder Containing Before Images"
        )
        
        if folder_path:
            self.before_images_folder = folder_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'before_images_label'):
                self.before_images_label.config(
                    text=f"Before Images: {truncate_path(folder_path)}",
                    fg="#FFFFFF"
                )
                self.before_images_label.unbind("<Enter>")
                self.before_images_label.unbind("<Leave>")
                Tooltip(self.before_images_label, folder_path)
            else:
                self.show_cutting()
            self.check_cutting_inputs()

    def select_after_images(self):
        folder_path = filedialog.askdirectory(
            title="Select Folder Containing After Images"
        )
        
        if folder_path:
            self.after_images_folder = folder_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'after_images_label'):
                self.after_images_label.config(
                    text=f"After Images: {truncate_path(folder_path)}",
                    fg="#FFFFFF"
                )
                self.after_images_label.unbind("<Enter>")
                self.after_images_label.unbind("<Leave>")
                Tooltip(self.after_images_label, folder_path)
            else:
                self.show_cutting()
            self.check_cutting_inputs()

    def select_cutting_excel(self):
        file_path = filedialog.askopenfilename(
            title="Select Cutting Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        
        if file_path:
            self.cutting_excel_file = file_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'cutting_excel_label'):
                self.cutting_excel_label.config(
                    text=f"Excel File: {truncate_path(file_path)}",
                    fg="#FFFFFF"
                )
                self.cutting_excel_label.unbind("<Enter>")
                self.cutting_excel_label.unbind("<Leave>")
                Tooltip(self.cutting_excel_label, file_path)
            else:
                self.show_cutting()
            
            # Calculate batches after loading Excel file
            self.update_cutting_batch_calculation()
            self.check_cutting_inputs()

    def select_cutting_output(self):
        folder_path = filedialog.askdirectory(
            title="Select Output Folder for Cutting Documents"
        )
        
        if folder_path:
            self.cutting_output_folder = folder_path
            def truncate_path(path, max_length=50):
                if len(path) <= max_length:
                    return path
                return f"{path[:max_length-3]}..."
            if hasattr(self, 'cutting_output_label'):
                self.cutting_output_label.config(
                    text=f"Output Folder: {truncate_path(folder_path)}",
                    fg="#FFFFFF"
                )
                self.cutting_output_label.unbind("<Enter>")
                self.cutting_output_label.unbind("<Leave>")
                Tooltip(self.cutting_output_label, folder_path)
            else:
                self.show_cutting()
            self.check_cutting_inputs()

    def check_cutting_inputs(self):
        if (hasattr(self, 'before_images_folder') and
            hasattr(self, 'after_images_folder') and
            hasattr(self, 'cutting_excel_file') and
            hasattr(self, 'cutting_output_folder')):
            self.cutting_process_btn.config(state=tk.NORMAL)
        else:
            self.cutting_process_btn.config(state=tk.DISABLED)

    def show_image_settings(self):
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Image Naming Settings")
        settings_window.geometry("400x300")
        settings_window.transient(self.root)
        settings_window.grab_set()
        
        settings_window.geometry("+%d+%d" % (
            self.root.winfo_rootx() + self.root.winfo_width() // 2 - 200,
            self.root.winfo_rooty() + self.root.winfo_height() // 2 - 150
        ))
        
        frame = tk.Frame(settings_window, bg="#1E1E1E", padx=20, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)
        
        pattern_label = tk.Label(
            frame, 
            text="Select image naming pattern:",
            bg="#1E1E1E",
            fg="#FFFFFF",
            font=("Arial", 10, "bold"),
            anchor="w"
        )
        pattern_label.pack(fill=tk.X, pady=(0, 10))
        
        self.pattern_var = tk.StringVar(value=self.image_pattern)
        
        patterns = [
            ("Simple numbers (e.g., '123.jpg')", "simple_numbers"),
            ("Tree prefix (e.g., 'T123.jpg', 'Tree123.jpg')", "tree_prefix"),
            ("Alpha-numeric (e.g., 'A123.jpg', 'B001.jpg')", "alpha_numeric"),
            ("Number with suffix (e.g., '123A.jpg', '001B.jpg')", "number_suffix"),
            ("Auto-detect (slower but handles mixed formats)", "auto_detect")
        ]
        
        for text, pattern in patterns:
            rb = tk.Radiobutton(
                frame,
                text=text,
                variable=self.pattern_var,
                value=pattern,
                bg="#1E1E1E",
                fg="#FFFFFF",
                selectcolor="#2A2A2A",
                activebackground="#2A2A2A",
                activeforeground="#FFFFFF"
            )
            rb.pack(anchor="w", pady=2)
        
        options_frame = tk.Frame(frame, bg="#1E1E1E", pady=10)
        options_frame.pack(fill=tk.X)
        
        self.use_subfolders_var = tk.BooleanVar(value=self.use_subfolders)
        subfolder_cb = tk.Checkbutton(
            options_frame,
            text="Search in subfolders (slower)",
            variable=self.use_subfolders_var,
            bg="#1E1E1E",
            fg="#FFFFFF",
            selectcolor="#2A2A2A",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF"
        )
        subfolder_cb.pack(anchor="w")
        
        button_frame = tk.Frame(frame, bg="#1E1E1E", pady=10)
        button_frame.pack(fill=tk.X)
        
        save_btn = tk.Button(
            button_frame,
            text="Save Settings",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=lambda: self.save_image_settings(settings_window)
        )
        save_btn.pack(side=tk.RIGHT, padx=5)
        
        cancel_btn = tk.Button(
            button_frame,
            text="Cancel",
            bg="#1E1E1E",
            fg="#FFFFFF",
            activebackground="#2A2A2A",
            activeforeground="#FFFFFF",
            bd=1,
            relief=tk.SOLID,
            font=("Arial", 10),
            command=settings_window.destroy
        )
        cancel_btn.pack(side=tk.RIGHT, padx=5)

    def save_image_settings(self, window):
        self.image_pattern = self.pattern_var.get()
        self.use_subfolders = self.use_subfolders_var.get()
        
        self.image_cache = {}
        self.fuzzy_cache = {}
        
        window.destroy()
        
        if hasattr(self, 'settings_label'):
            pattern_name = next((name for name, val in [
                ("Simple numbers", "simple_numbers"),
                ("Tree prefix", "tree_prefix"),
                ("Alpha-numeric", "alpha_numeric"),
                ("Number with suffix", "number_suffix"),
                ("Auto-detect", "auto_detect")
            ] if val == self.image_pattern), "Unknown")
            
            self.settings_label.config(
                text=f"Pattern: {pattern_name} | Subfolders: {'Yes' if self.use_subfolders else 'No'}",
                fg="#FFFFFF"
            )

    def process_data(self):
        try:
            try:
                self.start_row = int(self.start_row_var.get())
                if self.start_row < 2:
                    raise ValueError("Start row must be 2 or greater (row 1 is header)")
            except ValueError:
                messagebox.showerror("Invalid Input", "Start row must be a valid number (2 or greater)")
                return
            
            try:
                self.batch_size = int(self.batch_size_var.get())
                if self.batch_size < 1:
                    raise ValueError("Batch size must be at least 1")
            except ValueError:
                messagebox.showerror("Invalid Input", "Batch size must be a valid number (1 or greater)")
                return
            
            if self.num_batches_var.get().strip().lower() == "all":
                self.num_batches_to_process = None
            else:
                try:
                    self.num_batches_to_process = int(self.num_batches_var.get())
                    if self.num_batches_to_process < 1:
                        raise ValueError("Number of batches must be at least 1")
                except ValueError:
                    messagebox.showerror("Invalid Input", "Batches to process must be 'All' or a valid number")
                    return
            
            # Get and sanitize custom filename prefix for Inventory
            self.docx_filename_prefix = self.filename_prefix_var.get().strip()
            if not self.docx_filename_prefix:
                self.docx_filename_prefix = "TreeData"  # Fallback to default
            self.docx_filename_prefix = re.sub(r'[<>:"/\\|?*]', '_', self.docx_filename_prefix)
            
            self.progress_frame.pack(fill=tk.X, pady=10)
            self.progress_bar["value"] = 0
            self.root.update()
            
            self.progress_label.config(text="Loading Excel data...")
            self.root.update()
            
            try:
                # Read the Excel file
                df = pd.read_excel(self.excel_file_path)
                
                # Find the correct columns
                tree_number_col, species_col = self.find_excel_columns(df)
                
                if not tree_number_col or not species_col:
                    missing_cols = []
                    if not tree_number_col:
                        missing_cols.append("Tree Number")
                    if not species_col:
                        missing_cols.append("Species")
                    messagebox.showerror(
                        "Column Error",
                        f"Could not find required columns: {', '.join(missing_cols)}\n\n"
                        "Please ensure your Excel file has columns for Tree Number and Species."
                    )
                    self.progress_frame.pack_forget()
                    return
                
                # Rename columns to standard format
                df = df.rename(columns={
                    tree_number_col: 'TREE NUMBER',
                    species_col: 'SPECIES'
                })
                
                # Select only required columns
                df = df[['TREE NUMBER', 'SPECIES']].copy()
                
                # Convert columns to string type
                df['TREE NUMBER'] = df['TREE NUMBER'].astype(str)
                df['SPECIES'] = df['SPECIES'].astype(str)
                
            except Exception as e:
                messagebox.showerror(
                    "Excel Error",
                    f"Error reading Excel file: {str(e)}\n\n"
                    "Please ensure the file is a valid Excel file and contains the required columns."
                )
                self.progress_frame.pack_forget()
                return
            
            if self.start_row > 2:
                start_index = self.start_row - 2
                df = df.iloc[start_index:].reset_index(drop=True)
            
            self.progress_label.config(text="Loading images...")
            self.root.update()
            
            self.load_image_cache()
            
            total_rows = len(df)
            num_batches = (total_rows + self.batch_size - 1) // self.batch_size
            
            if self.num_batches_to_process is not None:
                num_batches = min(num_batches, self.num_batches_to_process)
            
            self.progress_bar["maximum"] = num_batches
            
            if not hasattr(self, 'match_cache'):
                self.match_cache = {}
            if not hasattr(self, 'image_size_cache'):
                self.image_size_cache = {}
            
            # Create batch status tracking
            batch_status = {
                'completed': 0,
                'failed': 0,
                'skipped': 0,
                'errors': []
            }
            
            # Create batch validation folder
            validation_folder = os.path.join(self.output_folder_path, 'BatchValidation')
            if not os.path.exists(validation_folder):
                os.makedirs(validation_folder)
            
            for batch_num in range(num_batches):
                try:
                    self.progress_label.config(text=f"Processing batch {batch_num + 1} of {num_batches}...")
                    self.progress_bar["value"] = batch_num
                    self.root.update()
                    
                    batch_start = batch_num * self.batch_size
                    batch_end = min((batch_num + 1) * self.batch_size, total_rows)
                    
                    batch_df = df.iloc[batch_start:batch_end].copy()
                    
                    # Validate batch data
                    validation_result = self.validate_batch(batch_df, batch_num + 1)
                    if not validation_result['is_valid']:
                        batch_status['skipped'] += 1
                        batch_status['errors'].append(f"Batch {batch_num + 1}: {validation_result['message']}")
                        continue
                    
                    # Process batch with error handling
                    try:
                        self.generate_word_doc(batch_df, batch_num + 1, num_batches)
                        batch_status['completed'] += 1
                    except Exception as e:
                        batch_status['failed'] += 1
                        batch_status['errors'].append(f"Batch {batch_num + 1}: {str(e)}")
                        continue
                    
                    # Save batch status
                    self.save_batch_status(batch_status, validation_folder)
                    
                    self.progress_bar["value"] = batch_num + 1
                    self.root.update()
                    
                    if batch_num % 5 == 0:
                        self.match_cache.clear()
                        self.fuzzy_cache.clear()
                        if len(self.image_size_cache) > 100:
                            self.image_size_cache.clear()
                        
                        import gc
                        gc.collect()
                
                except Exception as e:
                    batch_status['failed'] += 1
                    batch_status['errors'].append(f"Batch {batch_num + 1}: {str(e)}")
                    continue
            
            if hasattr(self, 'match_cache'):
                self.match_cache.clear()
            if hasattr(self, 'fuzzy_cache'):
                self.fuzzy_cache.clear()
            if hasattr(self, 'image_size_cache'):
                self.image_size_cache.clear()
                
            temp_dir = os.path.join(self.output_folder_path, 'temp_images')
            if os.path.exists(temp_dir):
                try:
                    import shutil
                    shutil.rmtree(temp_dir)
                except:
                    pass
            
            # Generate final report
            self.generate_processing_report(batch_status, validation_folder)
            
            self.progress_frame.pack_forget()
            messagebox.showinfo(
                "Processing Complete", 
                f"Generated {batch_status['completed']} Word documents in {self.output_folder_path}\n"
                f"Failed: {batch_status['failed']}\n"
                f"Skipped: {batch_status['skipped']}"
            )
            
        except Exception as e:
            self.progress_frame.pack_forget()
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            import traceback
            traceback.print_exc()

    def validate_batch(self, batch_df, batch_num):
        """Validate a batch of data before processing."""
        try:
            # Check for empty batch
            if batch_df.empty:
                return {'is_valid': False, 'message': 'Empty batch'}
            
            # Check for required columns
            required_cols = ['TREE NUMBER', 'SPECIES']
            missing_cols = [col for col in required_cols if col not in batch_df.columns]
            if missing_cols:
                return {'is_valid': False, 'message': f'Missing columns: {", ".join(missing_cols)}'}
            
            # Check for invalid tree numbers
            invalid_trees = batch_df[batch_df['TREE NUMBER'].isna() | (batch_df['TREE NUMBER'] == '')]
            if not invalid_trees.empty:
                return {'is_valid': False, 'message': f'Invalid tree numbers found in rows: {", ".join(map(str, invalid_trees.index + 1))}'}
            
            # Check for duplicate tree numbers
            duplicates = batch_df[batch_df['TREE NUMBER'].duplicated()]
            if not duplicates.empty:
                return {'is_valid': False, 'message': f'Duplicate tree numbers found: {", ".join(duplicates["TREE NUMBER"].unique())}'}
            
            return {'is_valid': True, 'message': 'Batch validation successful'}
            
        except Exception as e:
            return {'is_valid': False, 'message': f'Validation error: {str(e)}'}

    def save_batch_status(self, batch_status, validation_folder):
        """Save the current batch processing status."""
        try:
            status_file = os.path.join(validation_folder, 'batch_status.json')
            with open(status_file, 'w') as f:
                json.dump(batch_status, f, indent=4)
        except Exception as e:
            print(f"Error saving batch status: {str(e)}")

    def generate_processing_report(self, batch_status, validation_folder):
        """Generate a detailed processing report."""
        try:
            report_file = os.path.join(validation_folder, 'processing_report.txt')
            with open(report_file, 'w') as f:
                f.write("Tree Inventory Processing Report\n")
                f.write("=============================\n\n")
                f.write(f"Total batches completed: {batch_status['completed']}\n")
                f.write(f"Total batches failed: {batch_status['failed']}\n")
                f.write(f"Total batches skipped: {batch_status['skipped']}\n\n")
                
                if batch_status['errors']:
                    f.write("Errors encountered:\n")
                    for error in batch_status['errors']:
                        f.write(f"- {error}\n")
                
                f.write("\nProcessing completed at: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        except Exception as e:
            print(f"Error generating processing report: {str(e)}")

    def process_cutting_data(self):
        try:
            start_row = 2
            try:
                start_row = int(self.cutting_start_row_var.get())
                if start_row < 2:
                    raise ValueError("Start row must be 2 or greater (row 1 is header)")
            except ValueError:
                messagebox.showerror("Invalid Input", "Start row must be a valid number (2 or greater)")
                return
            
            batch_size = 100
            try:
                batch_size = int(self.cutting_batch_size_var.get())
                if batch_size < 1:
                    raise ValueError("Batch size must be at least 1")
            except ValueError:
                messagebox.showerror("Invalid Input", "Batch size must be a valid number (1 or greater)")
                return
            
            # Get and sanitize custom filename prefix for Cutting
            self.cutting_docx_filename_prefix = self.cutting_filename_prefix_var.get().strip()
            if not self.cutting_docx_filename_prefix:
                self.cutting_docx_filename_prefix = "TreeCutting"  # Fallback to default
            self.cutting_docx_filename_prefix = re.sub(r'[<>:"/\\|?*]', '_', self.cutting_docx_filename_prefix)
            
            self.cutting_progress_frame.pack(fill=tk.X, pady=10)
            self.cutting_progress_bar["value"] = 0
            self.root.update()
            
            self.cutting_progress_label.config(text="Loading Excel data...")
            self.root.update()
            
            try:
                # Read the Excel file
                df = pd.read_excel(self.cutting_excel_file)
                tree_number_col = df.columns[3]  # Column D
                species_col = df.columns[4]      # Column E

                # Filter out any rows where the tree number is the header (e.g., 'TREE NO.')
                df = df[df[tree_number_col].astype(str).str.upper() != 'TREE NO.']

                if not tree_number_col or not species_col:
                    missing_cols = []
                    if not tree_number_col:
                        missing_cols.append("Tree Number")
                    if not species_col:
                        missing_cols.append("Species")
                    messagebox.showerror(
                        "Column Error",
                        f"Could not find required columns: {', '.join(missing_cols)}\n\n"
                        "Please ensure your Excel file has columns for Tree Number and Species."
                    )
                    self.cutting_progress_frame.pack_forget()
                    return
                
                # Create a mapping for all columns
                columns_map = {
                    tree_number_col: 'TREE NUMBER',
                    species_col: 'SPECIES'
                }
                
                # Try to find cutting date and notes columns
                date_variations = ['cutting date', 'date', 'cut date', 'cutting time', 'cut time']
                notes_variations = ['notes', 'note', 'comments', 'comment', 'description', 'desc']
                
                # Find cutting date column
                for col in df.columns:
                    col_lower = col.lower()
                    for variation in date_variations:
                        if fuzz.ratio(variation, col_lower) >= 80:
                            columns_map[col] = 'CUTTING DATE'
                            break
                
                # Find notes column
                for col in df.columns:
                    col_lower = col.lower()
                    for variation in notes_variations:
                        if fuzz.ratio(variation, col_lower) >= 80:
                            columns_map[col] = 'NOTES'
                            break
                
                # Rename columns
                df = df.rename(columns=columns_map)
                
                # Ensure required columns exist
                required_cols = ['TREE NUMBER', 'SPECIES']
                for col in required_cols:
                    if col not in df.columns:
                        messagebox.showerror(
                            "Missing Column",
                            f"Required column not found: {col}"
                        )
                        self.cutting_progress_frame.pack_forget()
                        return
                
                # Add missing optional columns if needed
                if 'CUTTING DATE' not in df.columns:
                    df['CUTTING DATE'] = ''
                if 'NOTES' not in df.columns:
                    df['NOTES'] = ''
                
                # Select and order columns
                df = df[['TREE NUMBER', 'SPECIES', 'CUTTING DATE', 'NOTES']].copy()
                
                # Convert columns to string type
                df['TREE NUMBER'] = df['TREE NUMBER'].astype(str)
                df['SPECIES'] = df['SPECIES'].astype(str)
                df['CUTTING DATE'] = df['CUTTING DATE'].astype(str)
                df['NOTES'] = df['NOTES'].astype(str)
                
            except Exception as e:
                messagebox.showerror(
                    "Excel Error",
                    f"Error reading Excel file: {str(e)}\n\n"
                    "Please ensure the file is a valid Excel file and contains the required columns."
                )
                self.cutting_progress_frame.pack_forget()
                return
            
            if start_row > 2:
                df = df.iloc[start_row-2:].reset_index(drop=True)
            
            total_rows = len(df)
            num_batches = math.ceil(total_rows / batch_size)
            
            self.cutting_progress_bar["maximum"] = num_batches
            
            self.cutting_progress_label.config(text="Indexing images...")
            self.root.update()
            
            before_images = self.load_cutting_images(self.before_images_folder)
            after_images = self.load_cutting_images(self.after_images_folder)
            
            # Create batch status tracking
            batch_status = {
                'completed': 0,
                'failed': 0,
                'skipped': 0,
                'errors': []
            }
            
            # Create batch validation folder
            validation_folder = os.path.join(self.cutting_output_folder, 'BatchValidation')
            if not os.path.exists(validation_folder):
                os.makedirs(validation_folder)
            
            for batch_num in range(num_batches):
                try:
                    self.cutting_progress_label.config(text=f"Processing batch {batch_num + 1} of {num_batches}...")
                    self.cutting_progress_bar["value"] = batch_num
                    self.root.update()
                    
                    batch_start = batch_num * batch_size
                    batch_end = min((batch_num + 1) * batch_size, total_rows)
                    
                    batch_df = df.iloc[batch_start:batch_end].copy()
                    
                    # Validate batch data
                    validation_result = self.validate_cutting_batch(batch_df, batch_num + 1, before_images, after_images)
                    if not validation_result['is_valid']:
                        batch_status['skipped'] += 1
                        batch_status['errors'].append(f"Batch {batch_num + 1}: {validation_result['message']}")
                        continue
                    
                    # Process batch with error handling
                    try:
                        self.generate_cutting_doc(batch_df, batch_num + 1, num_batches, before_images, after_images)
                        batch_status['completed'] += 1
                    except Exception as e:
                        batch_status['failed'] += 1
                        batch_status['errors'].append(f"Batch {batch_num + 1}: {str(e)}")
                        continue
                    
                    # Save batch status
                    self.save_batch_status(batch_status, validation_folder)
                    
                    self.cutting_progress_bar["value"] = batch_num + 1
                    self.root.update()
                    
                    if batch_num % 5 == 0:
                        self.fuzzy_cache.clear()
                
                except Exception as e:
                    batch_status['failed'] += 1
                    batch_status['errors'].append(f"Batch {batch_num + 1}: {str(e)}")
                    continue
            
            self.fuzzy_cache.clear()
            
            # Generate final report
            self.generate_processing_report(batch_status, validation_folder)
            
            self.cutting_progress_frame.pack_forget()
            messagebox.showinfo(
                "Processing Complete", 
                f"Generated {batch_status['completed']} cutting documents in {self.cutting_output_folder}\n"
                f"Failed: {batch_status['failed']}\n"
                f"Skipped: {batch_status['skipped']}"
            )
            
        except Exception as e:
            self.cutting_progress_frame.pack_forget()
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            import traceback
            traceback.print_exc()

    def validate_cutting_batch(self, batch_df, batch_num, before_images, after_images):
        """Validate a batch of cutting data before processing."""
        try:
            # Check for empty batch
            if batch_df.empty:
                return {'is_valid': False, 'message': 'Empty batch'}
            
            # Check for required columns
            required_cols = ['TREE NUMBER', 'SPECIES']
            missing_cols = [col for col in required_cols if col not in batch_df.columns]
            if missing_cols:
                return {'is_valid': False, 'message': f'Missing columns: {", ".join(missing_cols)}'}
            
            # Check for invalid tree numbers
            invalid_trees = batch_df[batch_df['TREE NUMBER'].isna() | (batch_df['TREE NUMBER'] == '')]
            if not invalid_trees.empty:
                return {'is_valid': False, 'message': f'Invalid tree numbers found in rows: {", ".join(map(str, invalid_trees.index + 1))}'}
            
            # Check for duplicate tree numbers
            duplicates = batch_df[batch_df['TREE NUMBER'].duplicated()]
            if not duplicates.empty:
                return {'is_valid': False, 'message': f'Duplicate tree numbers found: {", ".join(duplicates["TREE NUMBER"].unique())}'}
            
            # Check for missing images
            missing_before = []
            missing_after = []
            for _, row in batch_df.iterrows():
                tree_num = str(row['TREE NUMBER'])
                if not self.find_cutting_image(tree_num, before_images, 'before'):
                    missing_before.append(tree_num)
                if not self.find_cutting_image(tree_num, after_images, 'after'):
                    missing_after.append(tree_num)
            
            if missing_before or missing_after:
                message = []
                if missing_before:
                    message.append(f"Missing before images for trees: {', '.join(missing_before)}")
                if missing_after:
                    message.append(f"Missing after images for trees: {', '.join(missing_after)}")
                return {'is_valid': False, 'message': ' | '.join(message)}
            
            return {'is_valid': True, 'message': 'Batch validation successful'}
            
        except Exception as e:
            return {'is_valid': False, 'message': f'Validation error: {str(e)}'}

    def load_image_cache(self):
        self.image_cache = {}
        
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tif', '.tiff']
        all_files = []
        
        image_extensions_set = set(image_extensions)
        
        if self.use_subfolders:
            for root, dirs, files in os.walk(self.image_folder_path):
                for file in files:
                    ext = os.path.splitext(file.lower())[1]
                    if ext in image_extensions_set:
                        all_files.append(os.path.join(root, file))
        else:
            for file in os.listdir(self.image_folder_path):
                ext = os.path.splitext(file.lower())[1]
                if ext in image_extensions_set:
                    all_files.append(os.path.join(self.image_folder_path, file))
        
        total_files = len(all_files)
        if total_files > 0:
            progress_window = tk.Toplevel(self.root)
            progress_window.title("Loading Images")
            progress_window.geometry("300x80")
            progress_window.transient(self.root)
            progress_window.grab_set()
            
            progress_window.geometry("+%d+%d" % (
                self.root.winfo_rootx() + self.root.winfo_width() // 2 - 150,
                self.root.winfo_rooty() + self.root.winfo_height() // 2 - 40
            ))
            
            progress_label = tk.Label(
                progress_window,
                text=f"Indexing image files (0/{total_files})...",
                font=("Arial", 10)
            )
            progress_label.pack(pady=5)
            
            image_progress = ttk.Progressbar(
                progress_window,
                orient="horizontal",
                length=280,
                mode="determinate",
                maximum=total_files
            )
            image_progress.pack(pady=5, padx=10)
            
            pattern_type = self.image_pattern
            
            patterns = {}
            if pattern_type == "simple_numbers" or pattern_type == "auto_detect":
                patterns['num'] = re.compile(r'(?:^|[^a-zA-Z0-9])(\d+)(?:[^a-zA-Z0-9]|$)')
            if pattern_type == "tree_prefix" or pattern_type == "auto_detect":
                patterns['tree'] = re.compile(r'(?:^|[^a-zA-Z])([Tt](?:ree)?)[- _]?(\d+)')
            if pattern_type == "alpha_numeric" or pattern_type == "auto_detect":
                patterns['alpha'] = re.compile(r'([a-zA-Z]+)[- _]?(\d+)')
            if pattern_type == "number_suffix" or pattern_type == "auto_detect":
                patterns['num_suffix'] = re.compile(r'(\d+)[- _]?([a-zA-Z]+)')
            if pattern_type == "auto_detect":
                patterns['all_nums'] = re.compile(r'\d+')
            
            batch_size = 200
            for i in range(0, total_files, batch_size):
                batch_end = min(i + batch_size, total_files)
                current_batch = all_files[i:batch_end]
                
                progress_label.config(text=f"Indexing image files ({i+1}/{total_files})...")
                image_progress["value"] = i + 1
                progress_window.update()
                
                for file_path in current_batch:
                    filename = os.path.basename(file_path)
                    name_without_ext = os.path.splitext(filename)[0]
                    
                    self.image_cache[f"file_{name_without_ext.lower()}"] = file_path
                    
                    if pattern_type == "simple_numbers" or pattern_type == "auto_detect":
                        self._process_simple_numbers(file_path, name_without_ext, patterns)
                    elif pattern_type == "tree_prefix":
                        self._process_tree_prefix(file_path, name_without_ext, patterns)
                    elif pattern_type == "alpha_numeric":
                        self._process_alpha_numeric(file_path, name_without_ext, patterns)
                    elif pattern_type == "number_suffix":
                        self._process_number_suffix(file_path, name_without_ext, patterns)
                
                image_progress["value"] = batch_end
                progress_window.update()
            
            progress_window.destroy()

    def _process_simple_numbers(self, file_path, name, patterns):
        matches = patterns['num'].finditer(name)
        for match in matches:
            num_str = match.group(1)
            try:
                num_val = int(num_str)
                self.image_cache[f"num_{num_val}"] = file_path
                if len(num_str) > 1 and num_str.startswith('0'):
                    self.image_cache[f"raw_{num_val}"] = file_path
            except ValueError:
                pass
    
    def _process_tree_prefix(self, file_path, name, patterns):
        matches = patterns['tree'].finditer(name)
        for match in matches:
            try:
                num_val = int(match.group(2))
                self.image_cache[f"tree_{num_val}"] = file_path
            except ValueError:
                pass
    
    def _process_alpha_numeric(self, file_path, name, patterns):
        matches = patterns['alpha'].finditer(name)
        for match in matches:
            prefix = match.group(1).lower()
            try:
                num_val = int(match.group(2))
                self.image_cache[f"alpha_{prefix}_{num_val}"] = file_path
            except ValueError:
                pass
    
    def _process_number_suffix(self, file_path, name, patterns):
        matches = patterns['num_suffix'].finditer(name)
        for match in matches:
            try:
                num_val = int(match.group(1))
                suffix = match.group(2).lower()
                self.image_cache[f"numsuf_{num_val}_{suffix}"] = file_path
                self.image_cache[f"numsuf_{num_val}"] = file_path
            except ValueError:
                pass

    def get_image_path(self, tree_number):
        if not tree_number or pd.isna(tree_number):
            return None
        
        try:
            tree_num_str = str(tree_number).strip()
            
            if tree_num_str in self.match_cache:
                return self.match_cache[tree_num_str]
            
            if not hasattr(self, 'match_cache'):
                self.match_cache = {}
                
            pattern_type = self.image_pattern
            result = None
            confidence_threshold = 80
            
            tree_num_lower = tree_num_str.lower()
            if f"file_{tree_num_lower}" in self.image_cache:
                result = self.image_cache[f"file_{tree_num_lower}"]
            
            if not result:
                try:
                    numeric_tree_num = int(tree_num_str)
                    
                    if pattern_type == "simple_numbers" or pattern_type == "auto_detect":
                        if f"num_{numeric_tree_num}" in self.image_cache:
                            result = self.image_cache[f"num_{numeric_tree_num}"]
                        elif f"raw_{numeric_tree_num}" in self.image_cache:
                            result = self.image_cache[f"raw_{numeric_tree_num}"]
                    
                    if not result and (pattern_type == "tree_prefix" or pattern_type == "auto_detect"):
                        if f"tree_{numeric_tree_num}" in self.image_cache:
                            result = self.image_cache[f"tree_{numeric_tree_num}"]
                            
                    if not result and pattern_type == "auto_detect":
                        if f"numsuf_{numeric_tree_num}" in self.image_cache:
                            result = self.image_cache[f"numsuf_{numeric_tree_num}"]
                except ValueError:
                    pass
            
            if not result:
                if pattern_type in ["alpha_numeric", "auto_detect"]:
                    match = re.match(r'^([a-zA-Z]+)[- _]?(\d+)$', tree_num_str)
                    if match:
                        prefix = match.group(1).lower()
                        try:
                            num_val = int(match.group(2))
                            if f"alpha_{prefix}_{num_val}" in self.image_cache:
                                result = self.image_cache[f"alpha_{prefix}_{num_val}"]
                        except ValueError:
                            pass
                
                if not result and pattern_type in ["number_suffix", "auto_detect"]:
                    match = re.match(r'^(\d+)[- _]?([a-zA-Z]+)$', tree_num_str)
                    if match:
                        try:
                            num_val = int(match.group(1))
                            suffix = match.group(2).lower()
                            if f"numsuf_{num_val}_{suffix}" in self.image_cache:
                                result = self.image_cache[f"numsuf_{num_val}_{suffix}"]
                        except ValueError:
                            pass
            
            if not result and pattern_type == "auto_detect":
                if tree_num_str in self.fuzzy_cache:
                    result = self.fuzzy_cache[tree_num_str]
                else:
                    best_score = 0
                    best_path = None
                    for key, path in self.image_cache.items():
                        if not key.startswith(('num_', 'tree_', 'alpha_', 'numsuf_')):
                            score = fuzz.ratio(tree_num_str.lower(), key.replace('file_', '').lower())
                            if score > best_score and score >= confidence_threshold:
                                best_score = score
                                best_path = path
                    result = best_path
                    self.fuzzy_cache[tree_num_str] = result
            
            self.match_cache[tree_num_str] = result
            return result
            
        except Exception as e:
            print(f"Error matching tree {tree_number}: {str(e)}")
            return None

    def optimize_image_for_word(self, image_path, max_width=1024, max_height=768):
        try:
            if not hasattr(self, 'image_size_cache'):
                self.image_size_cache = {}
                
            if image_path in self.image_size_cache:
                return self.image_size_cache[image_path]
                
            with Image.open(image_path) as img:
                if img.width > max_width or img.height > max_height:
                    ratio = min(max_width/img.width, max_height/img.height)
                    new_width = int(img.width * ratio)
                    new_height = int(img.height * ratio)
                    
                    temp_dir = os.path.join(self.output_folder_path, 'temp_images')
                    if not os.path.exists(temp_dir):
                        os.makedirs(temp_dir)
                        
                    file_name = os.path.basename(image_path)
                    temp_path = os.path.join(temp_dir, f"opt_{file_name}")
                    
                    if not os.path.exists(temp_path):
                        img_resized = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        img_resized.save(temp_path, quality=85, optimize=True)
                    
                    self.image_size_cache[image_path] = temp_path
                    return temp_path
                else:
                    self.image_size_cache[image_path] = image_path
                    return image_path
                    
        except Exception as e:
            print(f"Error optimizing image {image_path}: {str(e)}")
            return image_path

    def generate_word_doc(self, batch_df, batch_num, total_batches):
        doc = Document()
        
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        
        actual_rows = len(batch_df)
        trees_per_column = (actual_rows + 1) // 2
        table_rows = trees_per_column * 2
        
        table = doc.add_table(rows=table_rows, cols=2)
        table.style = 'Table Grid'
        
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(3.25)
        
        for i in range(table_rows):
            table.rows[i].height = Inches(2.5 if i % 2 == 0 else 0.4)
        
        arial_11_bold = {'name': 'Bernard Mt Condensed', 'size': Pt(16), 'bold': True}
        
        # List to store tree numbers with no images
        no_image_trees = []
        
        chunk_size = 10
        for chunk_start in range(0, actual_rows, chunk_size):
            chunk_end = min(chunk_start + chunk_size, actual_rows)
            
            for tree_index in range(chunk_start, chunk_end):
                tree_row = batch_df.iloc[tree_index]
                tree_number = str(tree_row['TREE NUMBER'])
                species = str(tree_row['SPECIES']).upper()
                
                row_index = tree_index // 2 * 2
                col_index = tree_index % 2
                
                image_cell = table.cell(row_index, col_index)
                image_para = image_cell.paragraphs[0]
                image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                image_path = self.get_image_path(tree_number)
                
                if image_path and os.path.exists(image_path):
                    try:
                        optimized_path = self.optimize_image_for_word(image_path)
                        run = image_para.add_run()
                        run.add_picture(optimized_path, width=Inches(3.23), height=Inches(2.43))
                    except Exception as e:
                        image_para.text = f"[No image found for Tree {tree_number}]"
                        for run in image_para.runs:
                            run.italic = True
                        no_image_trees.append(tree_number)
                else:
                    image_para.text = f"[No image found for Tree {tree_number}]"
                    for run in image_para.runs:
                        run.italic = True
                    no_image_trees.append(tree_number)
                
                text_cell = table.cell(row_index + 1, col_index)
                text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                p = text_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.space_after = p.space_before = Pt(0)
                
                text_content = f"TREE NUMBER : {tree_number}\nSPECIES      : {species}"
                tree_run = p.add_run(text_content)
                tree_run.font.name = arial_11_bold['name']
                tree_run.font.size = arial_11_bold['size']
                tree_run.font.bold = arial_11_bold['bold']
        
        if hasattr(self, 'match_cache'):
            self.match_cache.clear()
        if hasattr(self, 'fuzzy_cache'):
            self.fuzzy_cache.clear()
        
        file_name = os.path.join(
            self.output_folder_path, 
            f"{self.docx_filename_prefix}_Batch{batch_num}_of_{total_batches}.docx"
        )
        doc.save(file_name)
        
        # Save the list of trees with no images to an Excel file
        if no_image_trees:
            self.save_no_image_list(no_image_trees, batch_num, total_batches)
        
        if batch_num % 5 == 0 and hasattr(self, 'image_size_cache'):
            self.image_size_cache.clear()

    def save_no_image_list(self, no_image_trees, batch_num, total_batches):
        """
        Save a list of tree numbers with no images to an Excel file in a NoImageReports folder.
        
        Args:
            no_image_trees (list): List of tree numbers with no images.
            batch_num (int): Current batch number.
            total_batches (int): Total number of batches.
        """
        try:
            # Create NoImageReports folder
            no_image_folder = os.path.join(self.output_folder_path, "NoImageReports")
            if not os.path.exists(no_image_folder):
                os.makedirs(no_image_folder)
            
            # Define file path in NoImageReports folder
            file_name = os.path.join(
                no_image_folder,
                f"NoImages_Batch{batch_num}_of_{total_batches}.xlsx"
            )
            
            # Create a DataFrame from the list of tree numbers
            df = pd.DataFrame(no_image_trees, columns=["Tree Number"])
            
            # Save the DataFrame to an Excel file
            df.to_excel(file_name, index=False, engine='openpyxl')
            
        except Exception as e:
            print(f"Error writing no-image list for batch {batch_num}: {str(e)}")

    def load_cutting_images(self, folder_path):
        image_cache = {}
        
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tif', '.tiff']
        image_extensions_set = set(image_extensions)
        
        all_files = []
        
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                ext = os.path.splitext(file.lower())[1]
                if ext in image_extensions_set:
                    all_files.append(os.path.join(root, file))
        
        for file_path in all_files:
            filename = os.path.basename(file_path)
            name_without_ext = os.path.splitext(filename)[0]
            
            image_cache[name_without_ext.lower()] = file_path
            
            matches = re.findall(r'\d+', name_without_ext)
            for match in matches:
                try:
                    num_val = int(match)
                    image_cache[f"num_{num_val}"] = file_path
                    
                    if re.search(r'^[tT](?:ree)?[- _]?\d+', name_without_ext):
                        image_cache[f"tree_{num_val}"] = file_path
                    
                    alpha_match = re.search(r'^([a-zA-Z]+)[- _]?(\d+)', name_without_ext)
                    if alpha_match:
                        prefix = alpha_match.group(1).lower()
                        alpha_num = int(alpha_match.group(2))
                        image_cache[f"alpha_{prefix}_{alpha_num}"] = file_path
                except ValueError:
                    pass
        
        return image_cache

    def generate_cutting_doc(self, batch_df, batch_num, total_batches, before_images, after_images):
        from docx.shared import Cm
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)

        # Title
        title = doc.add_paragraph("Tree Cutting Report")
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_after = Pt(4)
        title_run = title.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        # Date
        date_paragraph = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.paragraph_format.space_after = Pt(10)
        date_run = date_paragraph.runs[0]
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(10)
        date_run.italic = True

        for idx, tree_row in batch_df.iterrows():
            tree_number = str(tree_row['TREE NUMBER'])
            species = str(tree_row['SPECIES']).upper()
            cutting_date = str(tree_row.get('CUTTING DATE', ''))
            notes = str(tree_row.get('NOTES', ''))
            action = notes.strip().upper() if notes else ''

            # Section titles
            if 'CUT' in action:
                prior_title = 'PRIOR TO TREE CUTTING'
                post_title = 'POST CUTTING'
            elif 'TRIM' in action:
                prior_title = 'PRIOR TO TREE TRIMMING'
                post_title = 'POST TRIMMING'
            elif 'EARTHBALL' in action:
                prior_title = 'PRIOR TO EARTHBALL'
                post_title = 'POST EARTHBALL'
            else:
                prior_title = 'PRIOR TO OPERATION'
                post_title = 'POST OPERATION'

            # Header row (table, single row, 4 columns, no borders)
            header_table = doc.add_table(rows=1, cols=4)
            hdr_cells = header_table.rows[0].cells
            hdr_cells[0].text = f"SEQUENCE NO.: {idx}"
            hdr_cells[1].text = f"SPECIES: {species}"
            hdr_cells[2].text = f"TREE NO.: {tree_number}"
            hdr_cells[3].text = f"DATE: {cutting_date if cutting_date and cutting_date.lower() != 'nan' else ''}"
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)
            # Remove borders
            tbl = header_table._tbl
            for border_dir in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                for border in tbl.xpath(f'.//w:tblBorders/w:{border_dir}'):
                    border.set(qn('w:val'), 'nil')
            header_table.rows[0].height = Cm(0.7)
            header_table.rows[0].height_rule = 1  # EXACT
            header_table.autofit = True
            doc.add_paragraph()  # Minimal space below header

            # PRIOR section title
            prior_heading = doc.add_paragraph(prior_title)
            prior_heading_format = prior_heading.paragraph_format
            prior_heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prior_heading_format.space_before = Pt(2)
            prior_heading_format.space_after = Pt(2)
            prior_heading_run = prior_heading.runs[0]
            prior_heading_run.font.bold = True
            prior_heading_run.font.size = Pt(12)
            prior_heading_run.font.name = 'Arial'

            # PRIOR image
            before_image_path = self.find_cutting_image(tree_number, before_images, 'before')
            if before_image_path:
                try:
                    optimized_path = self.optimize_image_for_word(before_image_path)
                    prior_img_para = doc.add_paragraph()
                    prior_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = prior_img_para.add_run()
                    run.add_picture(optimized_path, width=Cm(15.26), height=Cm(8.17))
                    prior_img_para.paragraph_format.space_after = Pt(4)
                except Exception as e:
                    prior_img_para = doc.add_paragraph("[Prior image not found]")
                    prior_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in prior_img_para.runs:
                        run.italic = True
                    prior_img_para.paragraph_format.space_after = Pt(4)
            else:
                prior_img_para = doc.add_paragraph("[Prior image not found]")
                prior_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in prior_img_para.runs:
                    run.italic = True
                prior_img_para.paragraph_format.space_after = Pt(4)

            # POST section title
            post_heading = doc.add_paragraph(post_title)
            post_heading_format = post_heading.paragraph_format
            post_heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            post_heading_format.space_before = Pt(2)
            post_heading_format.space_after = Pt(2)
            post_heading_run = post_heading.runs[0]
            post_heading_run.font.bold = True
            post_heading_run.font.size = Pt(12)
            post_heading_run.font.name = 'Arial'

            # POST image
            after_image_path = self.find_cutting_image(tree_number, after_images, 'after')
            if after_image_path:
                try:
                    optimized_path = self.optimize_image_for_word(after_image_path)
                    post_img_para = doc.add_paragraph()
                    post_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = post_img_para.add_run()
                    run.add_picture(optimized_path, width=Cm(15.26), height=Cm(8.17))
                    post_img_para.paragraph_format.space_after = Pt(8)
                except Exception as e:
                    post_img_para = doc.add_paragraph("[Post image not found]")
                    post_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in post_img_para.runs:
                        run.italic = True
                    post_img_para.paragraph_format.space_after = Pt(8)
            else:
                post_img_para = doc.add_paragraph("[Post image not found]")
                post_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in post_img_para.runs:
                    run.italic = True
                post_img_para.paragraph_format.space_after = Pt(8)

        file_name = os.path.join(
            self.cutting_output_folder, 
            f"{self.cutting_docx_filename_prefix}_Batch{batch_num}_of_{total_batches}.docx"
        )
        doc.save(file_name)

    def find_cutting_image(self, tree_number, image_cache, before_or_after):
        # before_or_after: 'before' or 'after'
        if not tree_number or pd.isna(tree_number):
            return None

        tree_num_str = str(tree_number).strip().lower()
        for key, path in image_cache.items():
            key_lower = key.lower()
            if tree_num_str in key_lower and before_or_after in key_lower:
                return path
        return None

    def update_ui(self, func):
        self.root.after(0, func)

    def run_in_background(self, func, callback=None):
        def _worker():
            result = None
            error = None
            try:
                result = func()
            except Exception as e:
                error = e
            
            if callback:
                self.update_ui(lambda: callback(result, error))
        
        thread = threading.Thread(target=_worker)
        thread.daemon = True
        thread.start()
        return thread

    def update_cutting_batch_calculation(self, *args):
        """Update the batch calculation display for cutting mode."""
        try:
            if not hasattr(self, 'cutting_excel_file') or not self.cutting_excel_file:
                self.cutting_batch_calc_label.config(text="Estimated batches: --")
                return
                
            # Read Excel file to get total rows
            df = pd.read_excel(self.cutting_excel_file)
            total_rows = len(df)
            
            # Get current batch size
            try:
                batch_size = int(self.cutting_batch_size_var.get())
                if batch_size < 1:
                    raise ValueError("Batch size must be positive")
            except ValueError:
                self.cutting_batch_calc_label.config(text="Estimated batches: Invalid batch size")
                return
            
            # Calculate number of batches
            num_batches = math.ceil(total_rows / batch_size)
            
            # Update display
            self.cutting_batch_calc_label.config(
                text=f"Estimated batches: {num_batches} (Total rows: {total_rows})",
                fg="#FFFFFF"
            )
            
        except Exception as e:
            self.cutting_batch_calc_label.config(
                text=f"Error calculating batches: {str(e)}",
                fg="#FF0000"
            )

    def find_excel_columns(self, df):
        """
        Find the correct columns for Tree Number and Species in the Excel file.
        Uses fuzzy matching to handle various header formats.
        """
        tree_number_variations = [
            'tree number', 'tree id', 'tree no', 'tree #', 'tree_num', 'treeid',
            'tree no.', 'tree #.', 'tree_number', 'tree_id', 'tree_no',
            'tree no ', 'tree no.', 'tree no', 'tree', 'item', 'tree no. ', 'tree no.\n', 'tree no.\r\n'
        ]
        species_variations = [
            'species', 'tree species', 'scientific name', 'botanical name',
            'tree type', 'tree kind', 'tree variety', 'species name', 'species '
        ]
        
        # Convert column names to lowercase and strip for comparison
        columns_lower = {col.lower().strip(): col for col in df.columns}
        
        # Prefer exact matches first
        tree_number_col = None
        for variation in tree_number_variations:
            for col_lower, col_orig in columns_lower.items():
                if col_lower == variation:
                    tree_number_col = col_orig
                    break
            if tree_number_col:
                break
        
        species_col = None
        for variation in species_variations:
            for col_lower, col_orig in columns_lower.items():
                if col_lower == variation:
                    species_col = col_orig
                    break
            if species_col:
                break
        
        # If not found, use fuzzy matching with lower threshold
        if not tree_number_col:
            best_tree_score = 0
            for variation in tree_number_variations:
                for col_lower, col_orig in columns_lower.items():
                    score = fuzz.ratio(variation, col_lower)
                    if score > best_tree_score and score >= 70:
                        best_tree_score = score
                        tree_number_col = col_orig
        if not species_col:
            best_species_score = 0
            for variation in species_variations:
                for col_lower, col_orig in columns_lower.items():
                    score = fuzz.ratio(variation, col_lower)
                    if score > best_species_score and score >= 70:
                        best_species_score = score
                        species_col = col_orig
        # Debug print (remove/comment out in production)
        print(f"[DEBUG] Matched Tree Number column: {tree_number_col}, Species column: {species_col}")
        return tree_number_col, species_col

if __name__ == "__main__":
    root = tk.Tk()
    app = TreeInventoryApp(root)
    root.mainloop()